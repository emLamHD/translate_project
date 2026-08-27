from __future__ import annotations

import base64
import json
import os
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from lxml import etree

from translator.models import TMEntry
from translator.translation.memory import source_hash


@pytest.fixture
def approved_entry() -> TMEntry:
    source = "Cân chính xác 5 g mẫu."
    return TMEntry(
        "vi",
        "en",
        source,
        "Accurately weigh 5 g of the sample.",
        source_hash(source),
        "human_approved",
        "human:1",
        "Reviewer",
        "2026-08-27T00:00:00Z",
        True,
    )


@pytest.fixture
def synthetic_docx(tmp_path: Path) -> Path:
    path = tmp_path / "synthetic.docx"
    doc = Document()
    doc.add_heading("PHẠM VI", level=1)
    doc.add_paragraph("Cân chính xác 5 g mẫu.")
    doc.add_paragraph("Đoạn chưa có bản dịch.")
    doc.add_paragraph("Mục đánh số", style="List Number")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Bảng thử nghiệm"
    table.cell(1, 0).text = "5 g"
    table.cell(1, 1).text = "DOC.TEST-01"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "HEADER"
    section.footer.paragraphs[0].text = "FOOTER"
    png = tmp_path / "pixel.png"
    png.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    visual = doc.add_paragraph("Nhãn kèm hình cần bỏ qua ")
    visual.add_run().add_picture(str(png))
    equation = doc.add_paragraph("Công thức cần bỏ qua ")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    math.append(math_run)
    equation._p.append(math)
    doc.save(path)
    # Add mocked chart and embedded object parts + relationships. They are not
    # rendered, but exercise byte/relationship preservation gates.
    with zipfile.ZipFile(path) as zf:
        members = {name: zf.read(name) for name in zf.namelist()}
    members["word/charts/chart1.xml"] = (
        b"<c:chartSpace xmlns:c='http://schemas.openxmlformats.org/drawingml/2006/chart'/>"
    )
    members["word/embeddings/object1.bin"] = b"DETERMINISTIC-EMBEDDED-OBJECT"
    rel_name = "word/_rels/document.xml.rels"
    root = etree.fromstring(members[rel_name])
    ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    for rid, typ, target in (
        (
            "rIdChartMock",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart",
            "charts/chart1.xml",
        ),
        (
            "rIdEmbedMock",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
            "embeddings/object1.bin",
        ),
    ):
        node = etree.SubElement(root, f"{{{ns}}}Relationship")
        node.set("Id", rid)
        node.set("Type", typ)
        node.set("Target", target)
    members[rel_name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
    temporary = path.with_suffix(".rewrite.docx")
    with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, payload in members.items():
            zf.writestr(name, payload)
    os.replace(temporary, path)
    return path


def write_tm(path: Path, entries: list[TMEntry]) -> Path:
    path.write_text(
        json.dumps({"schema_version": 1, "entries": [entry.to_dict() for entry in entries]}),
        encoding="utf-8",
    )
    return path
