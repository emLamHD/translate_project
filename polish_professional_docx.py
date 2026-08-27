from __future__ import annotations

import hashlib
import re
import zipfile
from collections import Counter
from copy import deepcopy
from pathlib import Path

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "GB2760-2024_song_ngu_Trung-Viet.docx"
SOURCE_PDF = ROOT / "source.pdf"
OUTPUT = ROOT / "GB2760-2024_song_ngu_Trung-Viet_chuyen_nghiep.docx"

CJK_RE = re.compile(r"[\u3400-\u9fff]")
CJK_STRIP_RE = re.compile(r"[^\u3400-\u9fff]+")

NAVY = "17365D"
NAVY_2 = "1F4E78"
TEAL = "2F6B73"
SLATE = "475569"
CHARCOAL = "273444"
MUTED = "64748B"
PALE_BLUE = "EAF2F8"
PALEST_BLUE = "F6F9FC"
LINE = "CBD5E1"
WHITE = "FFFFFF"


TOC_ROWS = [
    ("前言", "Lời nói đầu", "Ⅲ"),
    ("1 范围", "1 Phạm vi", "1"),
    ("2 术语和定义", "2 Thuật ngữ và định nghĩa", "1"),
    ("3 食品添加剂的使用原则", "3 Nguyên tắc sử dụng phụ gia thực phẩm", "1"),
    ("4 食品分类系统", "4 Hệ thống phân loại thực phẩm", "2"),
    ("5 食品添加剂的使用规定", "5 Quy định về sử dụng phụ gia thực phẩm", "2"),
    ("6 食品用香料", "6 Hương liệu dùng cho thực phẩm", "2"),
    ("7 食品工业用加工助剂", "7 Chất hỗ trợ chế biến dùng trong công nghiệp thực phẩm", "2"),
    ("8 食品添加剂的功能类别", "8 Nhóm chức năng của phụ gia thực phẩm", "2"),
    ("9 附录A 中食品添加剂使用规定索引", "9 Chỉ mục quy định sử dụng phụ gia thực phẩm trong Phụ lục A", "2"),
    ("10 营养强化剂", "10 Chất tăng cường dinh dưỡng", "3"),
    ("11 胶基糖果中基础剂物质", "11 Chất nền trong kẹo cao su", "3"),
    ("附录A 食品添加剂的使用规定", "Phụ lục A — Quy định về sử dụng phụ gia thực phẩm", "4"),
    ("附录B 食品用香料使用规定", "Phụ lục B — Quy định về sử dụng hương liệu dùng cho thực phẩm", "148"),
    ("附录C 食品工业用加工助剂使用规定", "Phụ lục C — Quy định về chất hỗ trợ chế biến dùng trong công nghiệp thực phẩm", "223"),
    ("附录D 食品添加剂功能类别", "Phụ lục D — Nhóm chức năng của phụ gia thực phẩm", "240"),
    ("附录E 食品分类系统", "Phụ lục E — Hệ thống phân loại thực phẩm", "241"),
    ("附录F 附录A 中食品添加剂使用规定索引", "Phụ lục F — Chỉ mục quy định sử dụng phụ gia thực phẩm trong Phụ lục A", "252"),
    ("表A.1 食品添加剂的允许使用品种、使用范围以及最大使用量或残留量", "Bảng A.1 — Phụ gia thực phẩm được phép sử dụng, phạm vi sử dụng và mức sử dụng tối đa hoặc lượng tồn dư", "5"),
    ("表A.2 表A.1中例外食品编号对应的食品类别", "Bảng A.2 — Loại thực phẩm tương ứng với mã thực phẩm ngoại lệ trong Bảng A.1", "146"),
    ("表B.1 不得添加食品用香料、香精的食品名单", "Bảng B.1 — Danh mục thực phẩm không được bổ sung hương liệu và hương tinh dùng cho thực phẩm", "149"),
    ("表B.2 允许使用的食品用天然香料名单", "Bảng B.2 — Danh mục hương liệu tự nhiên được phép sử dụng", "150"),
    ("表B.3 允许使用的食品用合成香料名单", "Bảng B.3 — Danh mục hương liệu tổng hợp được phép sử dụng", "165"),
    ("表C.1 可在各类食品加工过程中使用,残留量不需限定的加工助剂名单(不含酶制剂)", "Bảng C.1 — Chất hỗ trợ chế biến được sử dụng trong các quá trình chế biến thực phẩm, không cần quy định giới hạn lượng tồn dư (không gồm chế phẩm enzym)", "223"),
    ("表C.2 需要规定功能和使用范围的加工助剂名单(不含酶制剂)", "Bảng C.2 — Chất hỗ trợ chế biến cần quy định chức năng và phạm vi sử dụng (không gồm chế phẩm enzym)", "224"),
    ("表C.3 食品用酶制剂及其来源名单", "Bảng C.3 — Chế phẩm enzym dùng cho thực phẩm và nguồn gốc", "230"),
    ("表E.1 食品分类系统", "Bảng E.1 — Hệ thống phân loại thực phẩm", "241"),
]


STYLE_NAMES = {
    "cn": "Pro Chinese",
    "vi": "Pro Vietnamese",
    "heading_cn": "Pro Heading Chinese",
    "heading_vi": "Pro Heading Vietnamese",
    "title_cn": "Pro Title Chinese",
    "title_vi": "Pro Title Vietnamese",
    "entity_cn": "Pro Entity Chinese",
    "entity_vi": "Pro Entity Vietnamese",
    "meta": "Pro Metadata",
    "running": "Pro Running",
    "toc_cn": "Pro TOC Chinese",
    "toc_vi": "Pro TOC Vietnamese",
}


def cjk_chars(text: str) -> str:
    return CJK_STRIP_RE.sub("", text)


def has_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text))


def xml_hash(table) -> str:
    payload = etree.tostring(table._tbl, method="c14n", with_comments=False)
    return hashlib.sha256(payload).hexdigest()


def add_style(doc: Document, name: str, font: str, size: float, color: str, *, bold=False, italic=False):
    styles = doc.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)
    rfonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei" if "Chinese" in name else font)
    return style


def create_styles(doc: Document) -> None:
    add_style(doc, STYLE_NAMES["cn"], "SimSun", 9.8, CHARCOAL)
    add_style(doc, STYLE_NAMES["vi"], "Arial", 9.0, SLATE)
    add_style(doc, STYLE_NAMES["heading_cn"], "Microsoft YaHei", 11.5, NAVY, bold=True)
    add_style(doc, STYLE_NAMES["heading_vi"], "Arial", 9.7, TEAL, bold=True)
    add_style(doc, STYLE_NAMES["title_cn"], "Microsoft YaHei", 16.0, NAVY, bold=True)
    add_style(doc, STYLE_NAMES["title_vi"], "Arial", 11.5, SLATE, bold=True)
    add_style(doc, STYLE_NAMES["entity_cn"], "Microsoft YaHei", 10.7, NAVY_2, bold=True)
    add_style(doc, STYLE_NAMES["entity_vi"], "Arial", 9.3, TEAL, bold=True)
    add_style(doc, STYLE_NAMES["meta"], "Arial", 8.4, MUTED, italic=True)
    add_style(doc, STYLE_NAMES["running"], "Arial", 8.0, MUTED)
    add_style(doc, STYLE_NAMES["toc_cn"], "Microsoft YaHei", 9.5, NAVY_2, bold=True)
    add_style(doc, STYLE_NAMES["toc_vi"], "Arial", 8.3, SLATE)


def set_run_format(run, font: str, size: float, color: str, *, bold=False, italic=False, spacing: int | None = None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei" if font in ("Microsoft YaHei", "SimSun") else font)
    if spacing is not None:
        node = rpr.find(qn("w:spacing"))
        if node is None:
            node = OxmlElement("w:spacing")
            rpr.append(node)
        node.set(qn("w:val"), str(spacing))


def format_all_runs(paragraph, font: str, size: float, color: str, *, bold=False, italic=False, spacing=None):
    if not paragraph.runs and paragraph.text:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        set_run_format(run, font, size, color, bold=bold, italic=italic, spacing=spacing)


def remove_presentation_nodes(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for tag in ("w:pBdr", "w:shd"):
        node = ppr.find(qn(tag))
        if node is not None:
            ppr.remove(node)


def set_paragraph_border(paragraph, *, side="bottom", color=LINE, size=6, space=3, val="single") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = pbdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        pbdr.append(edge)
    edge.set(qn("w:val"), val)
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def reset_paragraph(paragraph, *, before=0, after=0, line=1.05, left=0, right=0, first=0):
    remove_presentation_nodes(paragraph)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)
    fmt.first_line_indent = Inches(first)
    fmt.keep_together = True
    fmt.widow_control = True


def set_cell_margins(cell, top=55, start=75, bottom=55, end=75):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, side="bottom", color=LINE, size=4, val="single"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    edge = borders.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        borders.append(edge)
    edge.set(qn("w:val"), val)
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:color"), color)


def shade_cell(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_page_border(doc: Document):
    sectpr = doc.sections[0]._sectPr
    old = sectpr.find(qn("w:pgBorders"))
    if old is not None:
        sectpr.remove(old)
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "20")
        edge.set(qn("w:color"), "D9E2F3")
        borders.append(edge)
    sectpr.append(borders)


def move_before(element, reference) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    reference.addprevious(element)


def add_toc_paragraph(doc: Document, text: str, style: str, reference, *, align, font, size, color, bold=False, after=0):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_format(run, font, size, color, bold=bold)
    move_before(p._p, reference)
    return p


def rebuild_toc(doc: Document, pdf: fitz.Document):
    raw_page_cjk = Counter(cjk_chars(pdf[1].get_text("text")))
    rebuilt_cjk = Counter(cjk_chars("目次" + "".join(row[0] for row in TOC_ROWS)))
    if raw_page_cjk != rebuilt_cjk:
        missing = sum((raw_page_cjk - rebuilt_cjk).values())
        extra = sum((rebuilt_cjk - raw_page_cjk).values())
        raise RuntimeError(f"Mục lục dựng lại lệch ký tự Hán: thiếu={missing}, thừa={extra}")

    page_breaks = [p._p for p in doc.paragraphs if p._p.xpath('.//w:br[@w:type="page"]')]
    if len(page_breaks) < 2:
        raise RuntimeError("Không tìm thấy đủ ngắt trang để dựng lại mục lục.")
    first_break, second_break = page_breaks[0], page_breaks[1]
    node = first_break.getnext()
    while node is not None and node is not second_break:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt

    add_toc_paragraph(
        doc, "GB2760—2024", STYLE_NAMES["running"], second_break,
        align=WD_ALIGN_PARAGRAPH.RIGHT, font="Arial", size=8.0, color=MUTED, after=5,
    )
    title_cn = add_toc_paragraph(
        doc, "目  次", STYLE_NAMES["title_cn"], second_break,
        align=WD_ALIGN_PARAGRAPH.CENTER, font="Microsoft YaHei", size=18, color=NAVY, bold=True, after=1,
    )
    set_paragraph_border(title_cn, color=NAVY_2, size=10, space=5)
    add_toc_paragraph(
        doc, "MỤC LỤC", STYLE_NAMES["title_vi"], second_break,
        align=WD_ALIGN_PARAGRAPH.CENTER, font="Arial", size=10.5, color=TEAL, bold=True, after=8,
    )

    table = doc.add_table(rows=len(TOC_ROWS), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    table.columns[1].width = Inches(0.55)
    for idx, (cn, vi, page_number) in enumerate(TOC_ROWS):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        left.width = Inches(6.55)
        right.width = Inches(0.55)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(left)
        set_cell_margins(right)
        left.text = ""
        p_cn = left.paragraphs[0]
        p_cn.style = STYLE_NAMES["toc_cn"]
        p_cn.paragraph_format.space_before = Pt(0)
        p_cn.paragraph_format.space_after = Pt(0)
        p_cn.paragraph_format.keep_with_next = True
        run = p_cn.add_run(cn)
        set_run_format(run, "Microsoft YaHei", 9.4, NAVY_2, bold=True)
        p_vi = left.add_paragraph(style=STYLE_NAMES["toc_vi"])
        p_vi.paragraph_format.space_before = Pt(0)
        p_vi.paragraph_format.space_after = Pt(0)
        run = p_vi.add_run(vi)
        set_run_format(run, "Arial", 8.2, SLATE)

        right.text = ""
        p_page = right.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_page.paragraph_format.space_before = Pt(0)
        p_page.paragraph_format.space_after = Pt(0)
        run = p_page.add_run(page_number)
        set_run_format(run, "Arial", 9.0, NAVY, bold=True)
        set_cell_border(left)
        set_cell_border(right)
        if idx in (12, 18):
            set_cell_border(left, side="top", color=NAVY_2, size=8)
            set_cell_border(right, side="top", color=NAVY_2, size=8)
        if idx % 2:
            shade_cell(left, PALEST_BLUE)
            shade_cell(right, PALEST_BLUE)
    move_before(table._tbl, second_break)

    footer = add_toc_paragraph(
        doc, "Ⅰ", STYLE_NAMES["running"], second_break,
        align=WD_ALIGN_PARAGRAPH.RIGHT, font="Arial", size=8, color=MUTED, after=0,
    )
    footer.paragraph_format.space_before = Pt(5)


def combine_split_section_numbers(doc: Document):
    paragraphs = doc.paragraphs
    remove = []
    for idx, p in enumerate(paragraphs[:-2]):
        if p.style.name != "Original Text":
            continue
        prefix = p.text.strip()
        if not re.fullmatch(r"(?:\d+|[A-F])\.", prefix):
            continue
        cn = paragraphs[idx + 1]
        vi = paragraphs[idx + 2]
        if cn.style.name != "Chinese Source" or vi.style.name != "Vietnamese Translation":
            continue
        if not re.match(r"^\s*\d+", cn.text):
            continue
        cn_bold = any(run.bold for run in cn.runs)
        vi_bold = any(run.bold for run in vi.runs)
        cn.text = prefix + cn.text.lstrip()
        vi.text = prefix + vi.text.lstrip()
        if cn.runs:
            cn.runs[0].bold = cn_bold
        if vi.runs:
            vi.runs[0].bold = vi_bold
        remove.append(p._p)
    for element in remove:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def pair_category(page: int, source, text: str) -> str:
    compact = re.sub(r"\s+", "", text)
    bold = any(bool(run.bold) for run in source.runs)
    sizes = [run.font.size.pt for run in source.runs if run.font.size]
    max_size = max(sizes) if sizes else 9.5
    if page == 1:
        return "cover"
    if max_size >= 13 or compact in ("前言", "目次") or re.fullmatch(r"附录[A-F]", compact):
        return "title"
    if page <= 7 and text.strip() in ("食品安全国家标准", "食品添加剂使用标准"):
        return "title"
    if bold:
        if page >= 8 and not re.match(r"^\s*\d{1,2}\s+", text):
            return "entity"
        return "heading"
    if re.match(r"^\s*(?:[a-z]\)|[—-]{2,}|\d+\))", text, re.IGNORECASE):
        return "list"
    if re.match(r"^\s*(?:CNS号|INS号|功能)", text):
        return "meta"
    return "normal"


def format_cover_pair(source, vietnamese, text: str):
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vietnamese.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact = re.sub(r"\s+", "", text)
    reset_paragraph(source, before=0, after=1, line=1.0)
    reset_paragraph(vietnamese, before=0, after=5, line=1.0)
    if compact == "中华人民共和国国家标准":
        source.paragraph_format.space_before = Pt(28)
        format_all_runs(source, "Microsoft YaHei", 17, NAVY, bold=True, spacing=28)
        format_all_runs(vietnamese, "Arial", 10.5, MUTED, bold=True)
    elif compact in ("食品安全国家标准", "食品添加剂使用标准"):
        size = 23 if compact == "食品添加剂使用标准" else 19
        format_all_runs(source, "Microsoft YaHei", size, NAVY, bold=True, spacing=12)
        format_all_runs(vietnamese, "Arial", 12.0, TEAL, bold=True)
        set_paragraph_shading(source, PALE_BLUE)
        set_paragraph_shading(vietnamese, PALE_BLUE)
    elif re.match(r"^\d{4}-\d{2}-\d{2}", compact):
        source.paragraph_format.space_before = Pt(5)
        format_all_runs(source, "Microsoft YaHei", 10.5, NAVY_2, bold=True)
        format_all_runs(vietnamese, "Arial", 8.8, SLATE)
    elif compact in ("发布",):
        format_all_runs(source, "Microsoft YaHei", 10.5, NAVY_2, bold=True)
        format_all_runs(vietnamese, "Arial", 8.8, SLATE)
    else:
        source.paragraph_format.space_before = Pt(5)
        format_all_runs(source, "Microsoft YaHei", 12.0, NAVY_2, bold=True)
        format_all_runs(vietnamese, "Arial", 9.2, SLATE)
    source.paragraph_format.keep_with_next = True
    vietnamese.paragraph_format.keep_together = True


def format_pair(page: int, source, vietnamese):
    text = source.text
    category = pair_category(page, source, text)
    source.paragraph_format.keep_with_next = True
    vietnamese.paragraph_format.keep_together = True
    if category == "cover":
        format_cover_pair(source, vietnamese, text)
        return
    if category == "title":
        source.style = STYLE_NAMES["title_cn"]
        vietnamese.style = STYLE_NAMES["title_vi"]
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vietnamese.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reset_paragraph(source, before=12, after=1, line=1.0)
        reset_paragraph(vietnamese, before=0, after=11, line=1.0)
        format_all_runs(source, "Microsoft YaHei", 16, NAVY, bold=True, spacing=12)
        format_all_runs(vietnamese, "Arial", 11.0, TEAL, bold=True)
        set_paragraph_border(source, color=NAVY_2, size=9, space=5)
        return
    if category == "heading":
        source.style = STYLE_NAMES["heading_cn"]
        vietnamese.style = STYLE_NAMES["heading_vi"]
        source.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vietnamese.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reset_paragraph(source, before=7, after=1, line=1.0)
        reset_paragraph(vietnamese, before=0, after=5, line=1.0, left=0.08)
        format_all_runs(source, "Microsoft YaHei", 11.5, NAVY, bold=True)
        format_all_runs(vietnamese, "Arial", 9.7, TEAL, bold=True)
        set_paragraph_border(source, color=LINE, size=6, space=3)
        return
    if category == "entity":
        source.style = STYLE_NAMES["entity_cn"]
        vietnamese.style = STYLE_NAMES["entity_vi"]
        source.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vietnamese.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reset_paragraph(source, before=5, after=0.5, line=1.0)
        reset_paragraph(vietnamese, before=0, after=2.5, line=1.0, left=0.12)
        format_all_runs(source, "Microsoft YaHei", 10.7, NAVY_2, bold=True)
        format_all_runs(vietnamese, "Arial", 9.3, TEAL, bold=True)
        set_paragraph_border(source, side="top", color=LINE, size=4, space=3)
        return
    if category == "meta":
        source.style = STYLE_NAMES["meta"]
        vietnamese.style = STYLE_NAMES["meta"]
        source.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vietnamese.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reset_paragraph(source, before=0, after=0, line=1.0, left=0.16)
        reset_paragraph(vietnamese, before=0, after=1, line=1.0, left=0.30)
        format_all_runs(source, "Microsoft YaHei", 8.6, MUTED)
        format_all_runs(vietnamese, "Arial", 8.1, MUTED, italic=True)
        return

    is_list = category == "list"
    source.style = STYLE_NAMES["cn"]
    vietnamese.style = STYLE_NAMES["vi"]
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vietnamese.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cn_left = 0.20 if is_list else 0
    vi_left = 0.42 if is_list else 0.24
    reset_paragraph(source, before=1.3, after=0, line=1.05, left=cn_left)
    reset_paragraph(vietnamese, before=0, after=3.0, line=1.05, left=vi_left, right=0.08)
    format_all_runs(source, "SimSun", 9.8, CHARCOAL)
    format_all_runs(vietnamese, "Arial", 9.0, SLATE)
    set_paragraph_border(vietnamese, side="left", color=LINE, size=5, space=5)


def next_is_page_break(paragraphs, idx: int) -> bool:
    return idx + 1 < len(paragraphs) and bool(paragraphs[idx + 1]._p.xpath('.//w:br[@w:type="page"]'))


def format_original(page: int, paragraph, paragraphs, idx: int):
    text = paragraph.text.strip()
    if not text:
        # Giữ nguyên vị trí ảnh; chỉ dọn khoảng trắng không cần thiết.
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if page == 1 else paragraph.alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return
    if text == "GB2760—2024":
        paragraph.style = STYLE_NAMES["running"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if page == 1 else WD_ALIGN_PARAGRAPH.RIGHT
        reset_paragraph(paragraph, before=4 if page == 1 else 0, after=10 if page == 1 else 4, line=1.0)
        format_all_runs(paragraph, "Arial", 17 if page == 1 else 8, NAVY if page == 1 else MUTED, bold=page == 1, spacing=12 if page == 1 else None)
        if page == 1:
            set_paragraph_border(paragraph, color=NAVY_2, size=8, space=5)
        else:
            set_paragraph_border(paragraph, color=LINE, size=4, space=2)
        return
    if next_is_page_break(paragraphs, idx) and re.fullmatch(r"(?:\d+|[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+)", text):
        paragraph.style = STYLE_NAMES["running"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        reset_paragraph(paragraph, before=3, after=0, line=1.0)
        format_all_runs(paragraph, "Arial", 8, MUTED)
        return
    if re.fullmatch(r"[…\.·]{5,}", text):
        paragraph.style = STYLE_NAMES["meta"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reset_paragraph(paragraph, before=0, after=0, line=1.0)
        format_all_runs(paragraph, "Arial", 3, LINE)
        return
    paragraph.style = STYLE_NAMES["meta"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reset_paragraph(paragraph, before=0, after=1.5, line=1.0, left=0.18)
    format_all_runs(paragraph, "Arial", 8.4, MUTED, italic=bool(re.search(r"[A-Za-z]", text)))


def style_body_outside_tables(doc: Document):
    paragraphs = doc.paragraphs
    page = 1
    idx = 0
    while idx < len(paragraphs):
        p = paragraphs[idx]
        if p._p.xpath('.//w:br[@w:type="page"]'):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            page += 1
            idx += 1
            continue
        if p.style.name == STYLE_NAMES["title_cn"] and idx + 1 < len(paragraphs):
            if paragraphs[idx + 1].style.name == STYLE_NAMES["title_vi"]:
                idx += 2
                continue
        if p.style.name.startswith("Pro "):
            idx += 1
            continue
        if p.style.name == "Chinese Source" and idx + 1 < len(paragraphs):
            vi = paragraphs[idx + 1]
            if vi.style.name == "Vietnamese Translation":
                format_pair(page, p, vi)
                idx += 2
                continue
        format_original(page, p, paragraphs, idx)
        idx += 1


def validate(original_hashes: list[str], pdf: fitz.Document):
    final = Document(OUTPUT)
    if len(final.tables) != len(original_hashes) + 1:
        raise RuntimeError(f"Số bảng không đúng: {len(final.tables)}")
    final_hashes = [xml_hash(t) for t in final.tables[1:]]
    if original_hashes != final_hashes:
        mismatches = sum(a != b for a, b in zip(original_hashes, final_hashes))
        raise RuntimeError(f"Có {mismatches} bảng nguồn bị thay đổi.")

    expected = Counter()
    for page in pdf:
        expected.update(cjk_chars(page.get_text("text")))
    actual = Counter(cjk_chars("".join(final.element.body.xpath(".//w:t/text()"))))
    if expected != actual:
        missing = sum((expected - actual).values())
        extra = sum((actual - expected).values())
        raise RuntimeError(f"Đối chiếu ký tự Hán thất bại: thiếu={missing}, thừa={extra}")

    pair_errors = 0
    cjk_paragraphs = 0
    for p in final.element.body.xpath(".//w:p"):
        text = "".join(p.xpath(".//w:t/text()"))
        if not has_cjk(text):
            continue
        cjk_paragraphs += 1
        nxt = p.getnext()
        if nxt is None or nxt.tag != qn("w:p"):
            pair_errors += 1
            continue
        next_text = "".join(nxt.xpath(".//w:t/text()"))
        if not next_text.strip() or has_cjk(next_text):
            pair_errors += 1
    if pair_errors:
        raise RuntimeError(f"Có {pair_errors} đoạn Trung không có đoạn Việt ngay dưới.")

    with zipfile.ZipFile(OUTPUT) as archive:
        corrupt = archive.testzip()
        if corrupt:
            raise RuntimeError(f"DOCX ZIP hỏng tại: {corrupt}")
    return {
        "tables_preserved": len(original_hashes),
        "toc_tables_added": 1,
        "cjk_paragraphs": cjk_paragraphs,
        "pair_errors": pair_errors,
        "cjk_characters": sum(expected.values()),
        "images": len(final.inline_shapes),
        "bytes": OUTPUT.stat().st_size,
    }


def main():
    doc = Document(SOURCE)
    original_hashes = [xml_hash(t) for t in doc.tables]
    pdf = fitz.open(SOURCE_PDF)
    create_styles(doc)
    set_page_border(doc)
    rebuild_toc(doc, pdf)
    combine_split_section_numbers(doc)
    style_body_outside_tables(doc)
    doc.core_properties.title = "GB 2760—2024 — Bản song ngữ Trung–Việt"
    doc.core_properties.subject = "Bản trình bày chuyên nghiệp"
    doc.save(OUTPUT)
    stats = validate(original_hashes, pdf)
    print(OUTPUT)
    for key, value in stats.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
