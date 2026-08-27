from __future__ import annotations

import io
import json
import re
import time
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter, OrderedDict
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_PDF = ROOT / "source.pdf"
OUTPUT_DOCX = ROOT / "GB2760-2024_song_ngu_Trung-Viet.docx"
CACHE_PATH = ROOT / "translation_cache.json"
REPORT_PATH = ROOT / "translation_report.json"

CJK_RE = re.compile(r"[\u3400-\u9fff]")
CJK_ONLY_RE = re.compile(r"[^\u3400-\u9fff]+")
SEP_RE = re.compile(r"\s*ZZZXSEP\s*(\d{6})\s*XXX\s*", re.IGNORECASE)

# Các mục xuất hiện thường xuyên được cố định để thuật ngữ pháp quy nhất quán.
EXACT_TRANSLATIONS = {
    "中华人民共和国国家标准": "Tiêu chuẩn quốc gia của Cộng hòa Nhân dân Trung Hoa",
    "食品安全国家标准": "Tiêu chuẩn quốc gia về an toàn thực phẩm",
    "食品添加剂使用标准": "Tiêu chuẩn sử dụng phụ gia thực phẩm",
    "中华人民共和国国家卫生健康委员会": "Ủy ban Y tế Quốc gia nước Cộng hòa Nhân dân Trung Hoa",
    "国家市场监督管理总局": "Tổng cục Quản lý thị trường Nhà nước",
    "发布": "Ban hành",
    "实施": "Thực hiện",
    "目  次": "Mục lục",
    "目 次": "Mục lục",
    "前  言": "Lời nói đầu",
    "前言": "Lời nói đầu",
    "范围": "Phạm vi",
    "术语和定义": "Thuật ngữ và định nghĩa",
    "食品添加剂": "Phụ gia thực phẩm",
    "食品分类号": "Mã phân loại thực phẩm",
    "食品名称": "Tên thực phẩm",
    "最大使用量": "Mức sử dụng tối đa",
    "最大使用量/(g/kg)": "Mức sử dụng tối đa/(g/kg)",
    "备注": "Ghi chú",
    "按生产需要适量使用": "Sử dụng với lượng thích hợp theo nhu cầu sản xuất",
    "序号": "Số thứ tự",
    "编码": "Mã số",
    "来源": "Nguồn",
    "供体": "Thể cho",
    "食品用香料": "Hương liệu dùng cho thực phẩm",
    "食品工业用加工助剂": "Chất hỗ trợ chế biến dùng trong công nghiệp thực phẩm",
    "营养强化剂": "Chất tăng cường dinh dưỡng",
    "食品分类系统": "Hệ thống phân loại thực phẩm",
}


def has_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text))


def cjk_chars(text: str) -> str:
    return CJK_ONLY_RE.sub("", text)


def rounded_bbox(bbox: Iterable[float]) -> list[float]:
    return [round(float(x), 3) for x in bbox]


def inside_bbox(cx: float, cy: float, bbox: Iterable[float], pad: float = 1.0) -> bool:
    x0, y0, x1, y1 = bbox
    return x0 - pad <= cx <= x1 + pad and y0 - pad <= cy <= y1 + pad


def make_line_model(line: dict[str, Any]) -> dict[str, Any]:
    spans = line.get("spans", [])
    text = "".join(span.get("text", "") for span in spans)
    sizes = [float(span.get("size", 9.7)) for span in spans if span.get("text")]
    fonts = [str(span.get("font", "")) for span in spans if span.get("text")]
    flags = [int(span.get("flags", 0)) for span in spans if span.get("text")]
    heading_font = any("HTK" in f or "XBSK" in f or "Bold" in f for f in fonts)
    return {
        "kind": "line",
        "bbox": rounded_bbox(line["bbox"]),
        "text": text,
        "size": round(max(sizes) if sizes else 9.7, 2),
        "bold": heading_font or any(flag & 16 for flag in flags),
    }


def table_model(table: Any) -> dict[str, Any]:
    data = table.extract()
    records: OrderedDict[tuple[float, ...], dict[str, Any]] = OrderedDict()
    for r, row in enumerate(table.rows):
        for c, bbox in enumerate(row.cells):
            if bbox is None:
                continue
            key = tuple(round(float(x), 3) for x in bbox)
            text = data[r][c] if r < len(data) and c < len(data[r]) else ""
            if key not in records:
                records[key] = {
                    "r": r,
                    "c": c,
                    "bbox": list(key),
                    "text": text or "",
                }
            elif not records[key]["text"] and text:
                records[key]["text"] = text
    return {
        "kind": "table",
        "bbox": rounded_bbox(table.bbox),
        "row_count": int(table.row_count),
        "col_count": int(table.col_count),
        "cells": list(records.values()),
        "images": [],
    }


def extract_structure(pdf: fitz.Document) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    pages: list[dict[str, Any]] = []
    all_pdf_cjk = Counter()
    all_model_cjk = Counter()
    table_total = 0
    table_row_total = 0
    table_cell_total = 0
    image_total = 0
    coverage_pages = 0

    for page_number, page in enumerate(pdf, start=1):
        tables_found = list(page.find_tables().tables)
        table_items = [table_model(t) for t in tables_found]
        table_bboxes = [item["bbox"] for item in table_items]

        items: list[dict[str, Any]] = []
        text_dict = page.get_text("dict", sort=True)
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                model = make_line_model(line)
                bbox = model["bbox"]
                cx = (bbox[0] + bbox[2]) / 2
                cy = (bbox[1] + bbox[3]) / 2
                if not any(inside_bbox(cx, cy, tb) for tb in table_bboxes):
                    items.append(model)

        items.extend(table_items)

        for image in page.get_image_info(xrefs=True):
            xref = int(image.get("xref", 0) or 0)
            if xref <= 0:
                continue
            bbox = rounded_bbox(image["bbox"])
            cx = (bbox[0] + bbox[2]) / 2
            cy = (bbox[1] + bbox[3]) / 2
            image_model = {"kind": "image", "bbox": bbox, "xref": xref}
            owner = None
            for tab in table_items:
                if inside_bbox(cx, cy, tab["bbox"], pad=0.2):
                    owner = tab
                    break
            if owner is None:
                items.append(image_model)
            else:
                owner["images"].append(image_model)
            image_total += 1

        kind_priority = {"image": 0, "line": 1, "table": 2}
        items.sort(key=lambda x: (round(x["bbox"][1], 2), kind_priority[x["kind"]], x["bbox"][0]))

        source_page_cjk = Counter(cjk_chars(page.get_text("text")))
        modeled_text = []
        for item in items:
            if item["kind"] == "line":
                modeled_text.append(item["text"])
            elif item["kind"] == "table":
                modeled_text.extend(cell["text"] for cell in item["cells"])
        model_page_cjk = Counter(cjk_chars("".join(modeled_text)))
        if source_page_cjk != model_page_cjk:
            missing = sum((source_page_cjk - model_page_cjk).values())
            extra = sum((model_page_cjk - source_page_cjk).values())
            raise RuntimeError(
                f"Trang {page_number}: mô hình trích xuất lệch ký tự Hán "
                f"(thiếu {missing}, thừa {extra})."
            )
        coverage_pages += 1
        all_pdf_cjk.update(source_page_cjk)
        all_model_cjk.update(model_page_cjk)

        table_total += len(table_items)
        table_row_total += sum(tab["row_count"] for tab in table_items)
        table_cell_total += sum(len(tab["cells"]) for tab in table_items)
        pages.append(
            {
                "number": page_number,
                "width": round(float(page.rect.width), 3),
                "height": round(float(page.rect.height), 3),
                "items": items,
            }
        )

    if all_pdf_cjk != all_model_cjk:
        raise RuntimeError("Đối chiếu tổng thể ký tự Hán của PDF không khớp mô hình.")

    stats = {
        "source_pages": len(pages),
        "coverage_pages": coverage_pages,
        "tables": table_total,
        "table_rows": table_row_total,
        "physical_table_cells": table_cell_total,
        "images": image_total,
        "source_cjk_characters": sum(all_pdf_cjk.values()),
    }
    return pages, stats


def iter_source_units(pages: list[dict[str, Any]]) -> Iterable[str]:
    for page in pages:
        for item in page["items"]:
            if item["kind"] == "line":
                text = item["text"].strip()
                if text and has_cjk(text):
                    yield text
            elif item["kind"] == "table":
                for cell in item["cells"]:
                    for raw_line in cell["text"].splitlines():
                        text = raw_line.strip()
                        if text and has_cjk(text):
                            yield text


def load_translation_cache() -> dict[str, str]:
    if not CACHE_PATH.exists():
        return {}
    with CACHE_PATH.open("r", encoding="utf-8") as stream:
        data = json.load(stream)
    return {str(k): str(v) for k, v in data.items() if str(v).strip()}


def save_translation_cache(cache: dict[str, str]) -> None:
    temp_path = CACHE_PATH.with_suffix(".json.tmp")
    with temp_path.open("w", encoding="utf-8", newline="\n") as stream:
        json.dump(cache, stream, ensure_ascii=False, indent=2, sort_keys=True)
    temp_path.replace(CACHE_PATH)


def google_translate_payload(payload: str, retries: int = 6) -> str:
    body = urllib.parse.urlencode(
        {"client": "gtx", "sl": "zh-CN", "tl": "vi", "dt": "t", "q": payload}
    ).encode("utf-8")
    request = urllib.request.Request(
        "https://translate.googleapis.com/translate_a/single",
        data=body,
        headers={"User-Agent": "Mozilla/5.0"},
    )
    last_error: Exception | None = None
    for attempt in range(retries):
        try:
            with urllib.request.urlopen(request, timeout=60) as response:
                result = json.loads(response.read().decode("utf-8"))
            if not result or not result[0]:
                raise RuntimeError("Dịch vụ dịch trả về dữ liệu rỗng.")
            return "".join(segment[0] or "" for segment in result[0])
        except (urllib.error.URLError, TimeoutError, OSError, ValueError, RuntimeError) as exc:
            last_error = exc
            if attempt + 1 < retries:
                time.sleep(min(2 ** attempt, 16))
    raise RuntimeError(f"Không thể gọi dịch vụ dịch sau {retries} lần: {last_error}")


def translate_individual(text: str) -> str:
    translated = google_translate_payload(text).strip()
    if not translated:
        raise RuntimeError(f"Bản dịch rỗng: {text!r}")
    return translated


def replace_case_insensitive(text: str, old: str, new: str) -> str:
    def replacement(match: re.Match[str]) -> str:
        value = new
        if match.group(0)[:1].isupper() and value:
            value = value[:1].upper() + value[1:]
        return value

    return re.sub(re.escape(old), replacement, text, flags=re.IGNORECASE)


def normalize_translation(source: str, translated: str) -> str:
    """Chuẩn hóa thuật ngữ pháp quy và bảo toàn mã hiệu quan trọng."""
    text = translated.strip()

    date_overrides = {
        "2024-02-08发布": "Ban hành ngày 2024-02-08",
        "2025-02-08实施": "Thực hiện từ ngày 2025-02-08",
    }
    if source in date_overrides:
        return date_overrides[source]

    if "加工助剂" in source:
        for variant in (
            "công cụ hỗ trợ xử lý",
            "công cụ hỗ trợ chế biến",
            "chất hỗ trợ xử lý",
            "trợ chất chế biến",
        ):
            text = replace_case_insensitive(text, variant, "chất hỗ trợ chế biến")

    has_flavoring = "食品用香料" in source
    has_essence = "食品用香精" in source or ("香精" in source and "食品用香料" in source)
    if "香料" in source:
        for variant in ("gia vị thực phẩm", "gia vị cho thực phẩm", "gia vị dùng cho thực phẩm"):
            text = replace_case_insensitive(text, variant, "hương liệu dùng cho thực phẩm")
        text = re.sub(r"(?i)\bgia vị\b", "hương liệu", text)
        text = replace_case_insensitive(text, "hương tự nhiên", "hương liệu tự nhiên")
        text = replace_case_insensitive(text, "hương tổng hợp", "hương liệu tổng hợp")

    if has_flavoring and not has_essence:
        text = replace_case_insensitive(text, "hương liệu thực phẩm", "hương liệu dùng cho thực phẩm")
        text = replace_case_insensitive(text, "hương vị cho thực phẩm", "hương liệu dùng cho thực phẩm")

    if has_essence:
        for variant in (
            "tinh chất thực phẩm",
            "hương liệu thực phẩm",
            "hương vị thực phẩm",
            "hương vị cho thực phẩm",
        ):
            text = replace_case_insensitive(text, variant, "hương tinh dùng cho thực phẩm")
        text = replace_case_insensitive(
            text,
            "phụ kiện tạo hương vị cho thực phẩm",
            "phụ liệu dùng cho hương tinh thực phẩm",
        )
    if "食品用香料、香精" in source:
        text = replace_case_insensitive(
            text,
            "hương liệu, hương liệu",
            "hương liệu dùng cho thực phẩm, hương tinh dùng cho thực phẩm",
        )
        text = replace_case_insensitive(
            text,
            "hương liệu, hương tinh dùng cho thực phẩm",
            "hương liệu dùng cho thực phẩm, hương tinh dùng cho thực phẩm",
        )
        text = replace_case_insensitive(
            text,
            "hương liệu và hương tinh dùng cho thực phẩm",
            "hương liệu dùng cho thực phẩm và hương tinh dùng cho thực phẩm",
        )

    if "酶制剂" in source:
        text = replace_case_insensitive(text, "chế phẩm enzyme", "chế phẩm enzym")

    if "最大使用量" in source:
        for variant in (
            "số lượng sử dụng tối đa",
            "công suất sử dụng tối đa",
            "liều lượng tối đa",
            "liều tối đa",
            "lượng sử dụng tối đa",
        ):
            text = replace_case_insensitive(text, variant, "mức sử dụng tối đa")

    residue_exact = {
        "残留量计": "Tính theo lượng tồn dư",
        "以残留量计": "Tính theo lượng tồn dư",
        "最大残留量": "Lượng tồn dư tối đa",
        "氧化硫残留量计": "Tính theo lượng tồn dư oxit lưu huỳnh",
    }
    if source in residue_exact:
        text = residue_exact[source]
    elif "残留量" in source:
        match = re.fullmatch(r"残留量([≤≥<>].+)", source)
        if match:
            text = f"Lượng tồn dư {match.group(1)}"
        else:
            for variant in ("lượng dư lượng", "lượng dư", "dư lượng"):
                text = replace_case_insensitive(text, variant, "lượng tồn dư")
            text = replace_case_insensitive(text, "số tiền còn lại", "lượng tồn dư")
            text = replace_case_insensitive(text, "máy đo cặn", "tính theo lượng tồn dư")
            text = replace_case_insensitive(text, "máy đo dư lượng", "tính theo lượng tồn dư")

    # Mã phân loại thực phẩm là mã hiệu, không được dịch thành ngày hoặc bị lược bỏ.
    food_codes = re.findall(r"(?<!\d)\d{2}(?:\.\d{2}){1,4}(?!\d)", source)
    missing_codes = []
    for code in food_codes:
        if code in text:
            continue
        parts = code.split(".")
        slash_variants = {"/".join(parts), "/".join(reversed(parts))}
        replaced = False
        for variant in slash_variants:
            if variant in text:
                text = text.replace(variant, code)
                replaced = True
                break
        if not replaced:
            missing_codes.append(code)
    if missing_codes:
        text = f"{'、'.join(missing_codes)} {text}"

    text = re.sub(r"[ \t]{2,}", " ", text).strip()
    return text


def translate_batches(units: list[str], cache: dict[str, str]) -> dict[str, str]:
    for source, vietnamese in EXACT_TRANSLATIONS.items():
        cache[source] = vietnamese

    pending = []
    seen = set()
    for unit in units:
        if unit not in cache and unit not in seen:
            pending.append(unit)
            seen.add(unit)

    batches: list[list[str]] = []
    current: list[str] = []
    current_size = 0
    max_payload_chars = 3400
    marker_cost = 24
    for text in pending:
        projected = current_size + len(text) + (marker_cost if current else 0)
        if current and projected > max_payload_chars:
            batches.append(current)
            current = []
            current_size = 0
        current.append(text)
        current_size += len(text) + (marker_cost if len(current) > 1 else 0)
    if current:
        batches.append(current)

    for batch_number, batch in enumerate(batches, start=1):
        if len(batch) == 1:
            cache[batch[0]] = translate_individual(batch[0])
        else:
            payload_parts = [batch[0]]
            for index, text in enumerate(batch[1:], start=1):
                payload_parts.append(f"\nZZZXSEP{index:06d}XXX\n")
                payload_parts.append(text)
            translated_payload = google_translate_payload("".join(payload_parts))
            split = SEP_RE.split(translated_payload)
            translations = [split[0].strip()]
            marker_ids = []
            for index in range(1, len(split), 2):
                marker_ids.append(int(split[index]))
                translations.append(split[index + 1].strip())
            expected_ids = list(range(1, len(batch)))
            if marker_ids != expected_ids or len(translations) != len(batch):
                translations = [translate_individual(text) for text in batch]
            for source, translated in zip(batch, translations):
                if not translated:
                    translated = translate_individual(source)
                cache[source] = translated
        save_translation_cache(cache)
        print(f"Dịch lô {batch_number}/{len(batches)}; cache={len(cache)}", flush=True)

    # Thử lại riêng các mục còn ký tự Hán; dịch riêng thường xử lý tốt hơn tên hóa chất ngắn.
    residual = [unit for unit in OrderedDict.fromkeys(units) if has_cjk(cache.get(unit, ""))]
    for source in residual:
        retried = translate_individual(source)
        if retried:
            cache[source] = retried

    for source in list(cache):
        cache[source] = normalize_translation(source, cache[source])
    for source, vietnamese in EXACT_TRANSLATIONS.items():
        cache[source] = vietnamese
    save_translation_cache(cache)

    missing = [unit for unit in units if not cache.get(unit, "").strip()]
    residual = [unit for unit in units if has_cjk(cache.get(unit, ""))]
    if missing:
        raise RuntimeError(f"Còn {len(missing)} đơn vị chưa có bản dịch.")
    if residual:
        sample = " | ".join(residual[:5])
        raise RuntimeError(f"Còn {len(residual)} bản dịch chứa ký tự Hán: {sample}")
    return cache


def set_run_font(run: Any, name: str, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(max(5.5, size_pt))
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "SimSun" if name != "Arial" else "Arial")


def set_paragraph_layout(paragraph: Any, *, keep_next: bool, after_pt: float = 0.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = 1.0
    fmt.keep_with_next = keep_next
    fmt.widow_control = True


def set_cell_margins(cell: Any, top: int = 30, start: int = 45, bottom: int = 30, end: int = 45) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell: Any, fill: str = "E7E6E6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cluster_coords(values: Iterable[float], tolerance: float = 0.8) -> list[float]:
    values = sorted(float(v) for v in values)
    groups: list[list[float]] = []
    for value in values:
        if not groups or abs(value - sum(groups[-1]) / len(groups[-1])) > tolerance:
            groups.append([value])
        else:
            groups[-1].append(value)
    return [sum(group) / len(group) for group in groups]


def nearest_index(coords: list[float], value: float) -> int:
    return min(range(len(coords)), key=lambda i: abs(coords[i] - value))


def append_styled_paragraph(
    container: Any,
    text: str,
    style_name: str,
    *,
    font_name: str,
    size_pt: float,
    bold: bool,
    italic: bool,
    keep_next: bool,
    after_pt: float,
    reuse_empty: bool = False,
) -> Any:
    paragraph = None
    if reuse_empty and hasattr(container, "paragraphs") and len(container.paragraphs) == 1:
        candidate = container.paragraphs[0]
        if not candidate.text and not candidate._p.xpath(".//w:drawing"):
            paragraph = candidate
            paragraph.style = style_name
    if paragraph is None:
        paragraph = container.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, bold=bold, italic=italic)
    set_paragraph_layout(paragraph, keep_next=keep_next, after_pt=after_pt)
    return paragraph


def apply_body_position(paragraph: Any, bbox: list[float], page_width: float, left_margin_pt: float) -> None:
    x0, _, x1, _ = bbox
    midpoint = (x0 + x1) / 2
    width = x1 - x0
    if abs(midpoint - page_width / 2) <= 15 and width < page_width * 0.78:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Pt(0)
    elif x0 > page_width * 0.64:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.left_indent = Pt(0)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(max(0, min(x0 - left_margin_pt, 110)))


def add_body_line(
    doc: Document,
    item: dict[str, Any],
    page_width: float,
    cache: dict[str, str],
    counters: Counter,
    left_margin_pt: float,
) -> None:
    text = item["text"]
    stripped = text.strip()
    cjk = bool(stripped and has_cjk(stripped))
    source_style = "Chinese Source" if cjk else "Original Text"
    source_size = float(item.get("size", 9.7))
    p_source = append_styled_paragraph(
        doc,
        text,
        source_style,
        font_name="SimSun" if cjk else "Arial",
        size_pt=source_size,
        bold=bool(item.get("bold", False)),
        italic=False,
        keep_next=cjk,
        after_pt=0,
    )
    apply_body_position(p_source, item["bbox"], page_width, left_margin_pt)
    if cjk:
        translation = cache[stripped]
        p_vi = append_styled_paragraph(
            doc,
            translation,
            "Vietnamese Translation",
            font_name="Arial",
            size_pt=max(7.0, source_size * 0.9),
            bold=bool(item.get("bold", False)),
            italic=True,
            keep_next=False,
            after_pt=1.2,
        )
        apply_body_position(p_vi, item["bbox"], page_width, left_margin_pt)
        counters["source_units"] += 1
        counters["translation_units"] += 1


def add_picture_to_paragraph(paragraph: Any, pdf: fitz.Document, xref: int, width_pt: float) -> None:
    image = pdf.extract_image(xref)
    stream = io.BytesIO(image["image"])
    run = paragraph.add_run()
    run.add_picture(stream, width=Pt(max(1.5, width_pt)))


def add_body_image(
    doc: Document,
    pdf: fitz.Document,
    item: dict[str, Any],
    page_width: float,
    left_margin_pt: float,
) -> None:
    paragraph = doc.add_paragraph(style="Original Text")
    set_paragraph_layout(paragraph, keep_next=False, after_pt=0)
    add_picture_to_paragraph(paragraph, pdf, item["xref"], item["bbox"][2] - item["bbox"][0])
    apply_body_position(paragraph, item["bbox"], page_width, left_margin_pt)


def add_cell_text(
    cell: Any,
    text: str,
    cache: dict[str, str],
    counters: Counter,
    *,
    header: bool,
) -> None:
    cell.text = ""
    first = True
    lines = text.splitlines() if text else []
    if not lines:
        return
    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            continue
        cjk = has_cjk(stripped)
        p_source = append_styled_paragraph(
            cell,
            raw_line,
            "Chinese Source" if cjk else "Original Text",
            font_name="SimSun" if cjk else "Arial",
            size_pt=7.1 if header else 6.7,
            bold=header,
            italic=False,
            keep_next=cjk,
            after_pt=0,
            reuse_empty=first,
        )
        first = False
        if header or (len(stripped) <= 14 and not re.search(r"[,，;；。]", stripped)):
            p_source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p_source.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if cjk:
            p_vi = append_styled_paragraph(
                cell,
                cache[stripped],
                "Vietnamese Translation",
                font_name="Arial",
                size_pt=6.2,
                bold=header,
                italic=True,
                keep_next=False,
                after_pt=0.7,
            )
            p_vi.alignment = p_source.alignment
            counters["source_units"] += 1
            counters["translation_units"] += 1


def add_table(
    doc: Document,
    pdf: fitz.Document,
    model: dict[str, Any],
    cache: dict[str, str],
    counters: Counter,
) -> None:
    rows = model["row_count"]
    cols = model["col_count"]
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    bbox = model["bbox"]
    x_values = [bbox[0], bbox[2]]
    y_values = [bbox[1], bbox[3]]
    for record in model["cells"]:
        x0, y0, x1, y1 = record["bbox"]
        x_values.extend((x0, x1))
        y_values.extend((y0, y1))
    x_coords = cluster_coords(x_values)
    y_coords = cluster_coords(y_values)
    if len(x_coords) != cols + 1:
        x_coords = [bbox[0] + (bbox[2] - bbox[0]) * i / cols for i in range(cols + 1)]
    if len(y_coords) != rows + 1:
        y_coords = [bbox[1] + (bbox[3] - bbox[1]) * i / rows for i in range(rows + 1)]

    total_pdf_width = max(1.0, x_coords[-1] - x_coords[0])
    column_widths = [
        Inches(7.2 * ((x_coords[i + 1] - x_coords[i]) / total_pdf_width))
        for i in range(cols)
    ]
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).width = column_widths[c]
            table.cell(r, c).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(table.cell(r, c))

    # Khôi phục các ô gộp dọc/ngang theo tọa độ đường kẻ bảng trong PDF.
    merged_keys = set()
    for record in model["cells"]:
        cell_bbox = record["bbox"]
        key = tuple(cell_bbox)
        if key in merged_keys:
            continue
        merged_keys.add(key)
        c0 = nearest_index(x_coords, cell_bbox[0])
        c1 = max(c0, nearest_index(x_coords, cell_bbox[2]) - 1)
        r0 = nearest_index(y_coords, cell_bbox[1])
        r1 = max(r0, nearest_index(y_coords, cell_bbox[3]) - 1)
        c0, c1 = min(c0, cols - 1), min(c1, cols - 1)
        r0, r1 = min(r0, rows - 1), min(r1, rows - 1)
        if r1 > r0 or c1 > c0:
            table.cell(r0, c0).merge(table.cell(r1, c1))

    for record in model["cells"]:
        r = min(int(record["r"]), rows - 1)
        c = min(int(record["c"]), cols - 1)
        cell = table.cell(r, c)
        add_cell_text(cell, record["text"], cache, counters, header=(r == 0))
        if r == 0:
            shade_cell(cell)

    if rows:
        mark_repeat_table_header(table.rows[0])

    # Các ảnh rất nhỏ là glyph/biểu tượng nhúng; giữ trong đúng ô chứa chúng.
    for image in model.get("images", []):
        cx = (image["bbox"][0] + image["bbox"][2]) / 2
        cy = (image["bbox"][1] + image["bbox"][3]) / 2
        c = min(max(nearest_index(x_coords, cx) - (1 if cx < x_coords[nearest_index(x_coords, cx)] else 0), 0), cols - 1)
        r = min(max(nearest_index(y_coords, cy) - (1 if cy < y_coords[nearest_index(y_coords, cy)] else 0), 0), rows - 1)
        cell = table.cell(r, c)
        paragraph = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
        add_picture_to_paragraph(paragraph, pdf, image["xref"], image["bbox"][2] - image["bbox"][0])


def configure_document(doc: Document) -> tuple[str, str, str]:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun")

    style_specs = (
        ("Chinese Source", "SimSun", False),
        ("Vietnamese Translation", "Arial", True),
        ("Original Text", "Arial", False),
    )
    style_ids = []
    for name, font_name, italic in style_specs:
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = font_name
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun" if name == "Chinese Source" else font_name)
        style_ids.append(style.style_id)

    doc.core_properties.title = "GB 2760—2024 Trung–Việt song ngữ"
    doc.core_properties.subject = "Bản song ngữ Trung–Việt"
    doc.core_properties.author = ""
    return tuple(style_ids)  # type: ignore[return-value]


def build_docx(
    pdf: fitz.Document,
    pages: list[dict[str, Any]],
    cache: dict[str, str],
) -> tuple[Counter, tuple[str, str, str]]:
    doc = Document()
    style_ids = configure_document(doc)
    counters = Counter()
    left_margin_pt = 1.4 / 2.54 * 72

    for page_index, page in enumerate(pages):
        for item in page["items"]:
            if item["kind"] == "line":
                add_body_line(doc, item, page["width"], cache, counters, left_margin_pt)
            elif item["kind"] == "image":
                add_body_image(doc, pdf, item, page["width"], left_margin_pt)
            elif item["kind"] == "table":
                add_table(doc, pdf, item, cache, counters)
        if page_index + 1 < len(pages):
            doc.add_page_break()
        if (page_index + 1) % 10 == 0 or page_index + 1 == len(pages):
            print(f"Dựng DOCX: {page_index + 1}/{len(pages)} trang nguồn", flush=True)

    doc.save(OUTPUT_DOCX)
    return counters, style_ids


def paragraph_text(element: Any) -> str:
    return "".join(element.xpath(".//w:t/text()"))


def validate_docx(
    expected_units: int,
    expected_cjk: Counter,
    build_counters: Counter,
    style_ids: tuple[str, str, str],
) -> dict[str, Any]:
    chinese_style_id, vietnamese_style_id, _ = style_ids
    check = Document(OUTPUT_DOCX)
    body = check.element.body
    source_nodes = body.xpath(f'.//w:p[w:pPr/w:pStyle[@w:val="{chinese_style_id}"]]')
    vietnamese_nodes = body.xpath(f'.//w:p[w:pPr/w:pStyle[@w:val="{vietnamese_style_id}"]]')

    adjacency_errors = 0
    for source in source_nodes:
        sibling = source.getnext()
        if sibling is None or sibling.tag != qn("w:p"):
            adjacency_errors += 1
            continue
        style_nodes = sibling.xpath("./w:pPr/w:pStyle/@w:val")
        if not style_nodes or style_nodes[0] != vietnamese_style_id:
            adjacency_errors += 1

    source_text = "".join(paragraph_text(node) for node in source_nodes)
    vietnamese_text = "".join(paragraph_text(node) for node in vietnamese_nodes)
    actual_cjk = Counter(cjk_chars(source_text))
    missing_cjk = sum((expected_cjk - actual_cjk).values())
    extra_cjk = sum((actual_cjk - expected_cjk).values())
    vietnamese_cjk = len(cjk_chars(vietnamese_text))

    checks = {
        "expected_source_units": expected_units,
        "inserted_source_units": build_counters["source_units"],
        "inserted_translation_units": build_counters["translation_units"],
        "docx_source_style_paragraphs": len(source_nodes),
        "docx_vietnamese_style_paragraphs": len(vietnamese_nodes),
        "pair_adjacency_errors": adjacency_errors,
        "missing_source_cjk_characters": missing_cjk,
        "extra_source_cjk_characters": extra_cjk,
        "cjk_characters_in_vietnamese": vietnamese_cjk,
        "output_bytes": OUTPUT_DOCX.stat().st_size,
    }
    required_zero = (
        adjacency_errors,
        missing_cjk,
        extra_cjk,
        vietnamese_cjk,
    )
    counts = (
        expected_units,
        build_counters["source_units"],
        build_counters["translation_units"],
        len(source_nodes),
        len(vietnamese_nodes),
    )
    if any(required_zero) or len(set(counts)) != 1:
        raise RuntimeError(f"Kiểm tra chống sót DOCX thất bại: {checks}")
    return checks


def main() -> None:
    started = time.time()
    pdf = fitz.open(SOURCE_PDF)
    print(f"Đọc PDF: {pdf.page_count} trang", flush=True)
    pages, stats = extract_structure(pdf)
    units = list(iter_source_units(pages))
    expected_cjk = Counter()
    for page in pdf:
        expected_cjk.update(cjk_chars(page.get_text("text")))
    stats["source_translation_unit_occurrences"] = len(units)
    stats["unique_source_translation_units"] = len(set(units))
    print(
        f"Trích xuất đủ {stats['coverage_pages']}/{stats['source_pages']} trang; "
        f"đơn vị cần dịch={len(units)}, duy nhất={len(set(units))}",
        flush=True,
    )

    cache = translate_batches(units, load_translation_cache())
    build_counters, style_ids = build_docx(pdf, pages, cache)
    checks = validate_docx(len(units), expected_cjk, build_counters, style_ids)

    report = {
        "status": "PASS",
        **stats,
        **checks,
        "elapsed_seconds": round(time.time() - started, 1),
        "output_file": OUTPUT_DOCX.name,
    }
    with REPORT_PATH.open("w", encoding="utf-8", newline="\n") as stream:
        json.dump(report, stream, ensure_ascii=False, indent=2)
    print(json.dumps(report, ensure_ascii=False, indent=2), flush=True)


if __name__ == "__main__":
    main()
